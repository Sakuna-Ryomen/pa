import streamlit as st
from docx import Document
import os
import io
import cohere
import logging
from dotenv import load_dotenv
from pyngrok import ngrok


# Load environment variables from .env file
load_dotenv()

# Get Cohere API key
api_key = os.getenv("COHERE_API_KEY")


# Setup logging
logging.basicConfig(level=logging.INFO)

# Valid sections
valid_sections = ["Introduction", "Problem Statement", "Objective", "Methodology", "Literature Review", "Future Scope"]


def generate_content(section, domain, title, api_key):
    """Generate content using Cohere Command R+ API"""
    if section not in valid_sections:
        return f"Invalid section: {section}"

    prompt = f"""Write a comprehensive {section} for a final year {domain} project titled '{title}' in a formal academic tone. 

The content should be:
- Professional and detailed
- Well-structured with clear paragraphs
- Academically appropriate for a final year project
- Approximately 200-300 words
- Include specific technical details relevant to {domain}

Please provide only the content for the {section} section, without any additional commentary."""

    try:
        co = cohere.Client(api_key)
        response = co.generate(
            model='command-r-plus',
            prompt=prompt,
            max_tokens=500,
            temperature=0.7,
            k=0,
            stop_sequences=[],
            return_likelihoods='NONE'
        )
        content = response.generations[0].text.strip()

        if len(content) < 50:
            return generate_fallback_content(section, domain, title)
        return content
    except Exception as e:
        logging.exception("Cohere API failed.")
        st.error(f"⚠️ Error generating content: {str(e)}")
        return generate_fallback_content(section, domain, title)


def generate_fallback_content(section, domain, title):
    """Generate fallback academic content when AI fails"""
    templates = {
        "Introduction": f"""The {title} project represents a significant advancement in the field of {domain}. This project aims to address current challenges and provide innovative solutions that can benefit both academic research and practical applications.

In today's rapidly evolving technological landscape, {domain} continues to play a crucial role in solving complex problems. This project focuses on developing comprehensive solutions that leverage modern technologies and methodologies.

The scope of this project encompasses various aspects of {domain}, including theoretical foundations, practical implementations, and real-world applications. Through systematic research and development, this project seeks to contribute meaningful insights to the field.""",

        "Problem Statement": f"""The current state of {domain} faces several challenges that require immediate attention and innovative solutions. Traditional approaches to {title} have limitations that hinder optimal performance and scalability.

Key issues identified include:
1. Limited efficiency in existing systems
2. Lack of comprehensive solutions
3. Need for improved user experience
4. Scalability concerns in current implementations

This project addresses these challenges by proposing a novel approach that combines theoretical knowledge with practical implementation strategies.""",

        "Objective": f"""The primary objective of this {title} project is to develop a comprehensive solution in the {domain} domain that addresses current limitations and provides enhanced functionality.

Specific objectives include:
1. To analyze existing systems and identify improvement opportunities
2. To design and implement an efficient {domain} solution
3. To evaluate performance and validate the proposed approach
4. To provide documentation and guidelines for future development

These objectives will be achieved through systematic research, careful planning, and rigorous testing procedures.""",

        "Methodology": f"""The methodology for this {title} project follows a systematic approach that ensures comprehensive coverage of all project aspects.

Phase 1: Research and Analysis
- Literature review of existing {domain} solutions
- Identification of current challenges and limitations
- Analysis of user requirements and expectations

Phase 2: Design and Development
- System architecture design
- Implementation of core functionalities
- Integration of {domain} technologies

Phase 3: Testing and Validation
- Performance testing and optimization
- User acceptance testing
- Documentation and deployment""",

        "Literature Review": f"""The literature review for this {title} project examines current research and developments in the {domain} field. Various studies have been conducted to understand the theoretical foundations and practical applications.

Recent research in {domain} has focused on improving efficiency, scalability, and user experience. Several authors have proposed different approaches to address current challenges.

Key findings from the literature include:
1. The importance of comprehensive system design
2. The need for efficient implementation strategies
3. The significance of user-centered approaches
4. The value of continuous improvement and optimization

This project builds upon existing knowledge while introducing novel concepts and methodologies.""",

        "Future Scope": f"""The {title} project provides a foundation for future developments in the {domain} field. Several opportunities exist for extending and improving the current solution.

Potential future enhancements include:
1. Integration with emerging technologies
2. Expansion of functionality and features
3. Performance optimization and scalability improvements
4. Development of mobile and web-based interfaces

The project also opens possibilities for research in related areas such as:
- Advanced algorithms and techniques
- Cross-platform compatibility
- Enhanced user experience design
- Integration with other {domain} systems

These future developments will contribute to the continuous evolution of {domain} solutions."""
    }
    return templates.get(section, f"Content for {section} section of {title} project in {domain} domain.")

# Launch ngrok tunnel to Streamlit default port 8501
# Safe ngrok tunnel creation — only once
if "ngrok_tunnel" not in st.session_state:
    try:
        public_url = ngrok.connect(8501, bind_tls=True)
        st.session_state["ngrok_tunnel"] = public_url
        st.sidebar.success("✅ Ngrok tunnel created")
    except Exception as e:
        st.sidebar.error(f"❌ Ngrok Error: {e}")
        st.stop()
else:
    public_url = st.session_state["ngrok_tunnel"]

st.sidebar.markdown(f"🌐 **Public Link:** [Open App]({public_url})")

# === Streamlit App ===
st.set_page_config(page_title="Blackbook Content Generator", layout="wide")
st.title("📚 Blackbook Content Generator using AI")

# Sidebar - About
st.sidebar.title("About")
st.sidebar.info("""
This app uses Cohere's Command R+ model to generate academic content 
for different sections of a project blackbook.

**Features:**
- Generate section-wise or full content
- Download as DOCX file
- Academic language
- Multiple domains supported
- Powered by Cohere Command R+ (FREE tier available)
""")

# Sidebar - API Key
st.sidebar.title("Configuration")
st.sidebar.markdown("""
**Steps to get FREE API key:**
1. Visit [cohere.ai](https://cohere.ai)
2. Create a free account
3. Go to API keys section
4. Generate a new API key
5. Paste below 👇
""")

api_key = os.getenv("COHERE_API_KEY", "")
if api_key:
    st.sidebar.success("✅ API Key loaded from .env")
else:
    st.sidebar.error("⚠️ API Key not found. Please add it to .env file.")


# Main content
st.markdown("### Project Details")
col1, col2 = st.columns(2)

with col1:
    project_title = st.text_input("📝 Enter the Project Title", placeholder="e.g., Smart Campus Surveillance using AI")
    domain = st.selectbox("🎯 Select Project Domain", ["AI", "IoT", "Web", "Data Science", "Cybersecurity"])

with col2:
    section = st.selectbox("📋 Select Section to Generate", valid_sections)
    generate_all = st.checkbox("📚 Generate Full Blackbook (All Sections)")

st.markdown("---")

# Generate content
if st.button("🚀 Generate Content", type="primary"):
    if not api_key:
        st.error("⚠️ Please enter your Cohere API key.")
    elif not project_title:
        st.error("⚠️ Please enter the project title.")
    else:
        if generate_all:
            with st.spinner("Generating full blackbook..."):
                doc = Document()
                doc.add_heading(project_title, 0)
                for sec in valid_sections:
                    content = generate_content(sec, domain, project_title, api_key)
                    doc.add_heading(sec, level=1)
                    doc.add_paragraph(content)

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.success("✅ Full Blackbook Generated!")
                st.download_button(
                    label="📥 Download Full Blackbook as DOCX",
                    data=buffer.getvalue(),
                    file_name=f"{project_title.replace(' ', '_')}_Blackbook.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            with st.spinner("Generating content..."):
                content = generate_content(section, domain, project_title, api_key)
                st.success("✅ Section generated successfully!")
                st.markdown("### Generated Content")
                st.text_area("📄 Output", content, height=400)

                doc = Document()
                doc.add_heading(f'{project_title} - {section}', 0)
                doc.add_paragraph(content)

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label="📥 Download Section as DOCX",
                    data=buffer.getvalue(),
                    file_name=f"{project_title.replace(' ', '_')}_{section}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# Footer
st.markdown("---")
st.markdown("<p style='text-align: center; color: gray;'>Made with ❤️ using Streamlit and Cohere Command R+</p>",
            unsafe_allow_html=True)
